"""Génère la fiche A4 des tables et sa source DOCX modifiable.

La mise en page reprend la fiche originale validée : logo et titre en haut,
QR code à droite, consigne courte, puis neuf cartes colorées. Les textes qui
seront le plus souvent adaptés sont regroupés juste sous les imports.
"""

from pathlib import Path
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.graphics import renderSVG
from reportlab.graphics.barcode.qr import QrCodeWidget
from reportlab.graphics.shapes import Drawing


TITLE_LINE_1 = "MES TABLES"
TITLE_LINE_2 = "DE MULTIPLICATION"
GUIDE_TITLE = "POUR COMMENCER SUR L’APPLICATION"
GUIDE_TEXT = "Suis le parcours : J’apprends  →  Je m’entraîne  →  Je deviens expert."
APP_URL = "https://mathsgo.re/outils/calcul_mental/defi_tables.html"
FOOTER_TEXT = "mathsgo.re  ·  CC BY-NC-SA 4.0"
TABLES = range(2, 11)
MULTIPLIERS = range(1, 11)

ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DOCX = ROOT / "_sources" / "defi-tables" / "fiche_tables_multiplication.docx"
WORK_DIR = ROOT / "tmp" / "defi-tables-fiche"
QR_PATH = WORK_DIR / "qr-defi-tables.png"

INK = "123452"
BLUE = "0B57B7"
TEAL = "087F75"
ORANGE = "E66D0A"
MUTED = "506B7D"
LINE = "CFE2ED"
PALE_TEAL = "E7F7F4"
WHITE = "FFFFFF"
TINTS = {BLUE: "EEF5FB", TEAL: "E8F6F3", ORANGE: "FFF1E5"}
FONT = "Arial"


def set_run_font(run, size, color=INK, bold=False, name=FONT):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), name)
    r_pr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_borders(cell, *, top=None, start=None, bottom=None, end=None, size=10):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, color in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil" if color is None else "single")
        if color is not None:
            border.set(qn("w:sz"), str(size))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr

    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")

    table_indent = tbl_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        tbl_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, value in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(value))
            tc_w.set(qn("w:type"), "dxa")


def set_row_height(row, centimetres):
    row.height = Cm(centimetres)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def add_qr(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    widget = QrCodeWidget(APP_URL)
    x0, y0, x1, y1 = widget.getBounds()
    width = x1 - x0
    height = y1 - y0
    target = 420
    drawing = Drawing(target, target, transform=[target / width, 0, 0, target / height, 0, 0])
    drawing.add(widget)
    svg_path = path.with_suffix(".svg")
    renderSVG.drawToFile(drawing, str(svg_path))
    subprocess.run(
        [
            "inkscape",
            str(svg_path),
            "--export-type=png",
            f"--export-filename={path}",
            "--export-width=420",
            "--export-height=420",
        ],
        check=True,
    )


def add_card(cell, table_number, accent):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=45, start=55, bottom=45, end=55)
    set_cell_borders(cell)
    clear_paragraph(cell.paragraphs[0])

    card = cell.add_table(rows=11, cols=1)
    set_table_geometry(card, [3160])
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, row in enumerate(card.rows):
        set_row_height(row, 0.56 if index else 0.88)
        card_cell = row.cells[0]
        card_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(card_cell, top=0, start=40, bottom=0, end=40)
        set_cell_borders(
            card_cell,
            top=LINE if index == 0 else None,
            start=LINE,
            bottom=LINE if index == 10 else None,
            end=LINE,
            size=9,
        )
        paragraph = clear_paragraph(card_cell.paragraphs[0])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if index == 0:
            set_cell_fill(card_cell, accent)
            run = paragraph.add_run(f"TABLE DE {table_number}")
            set_run_font(run, 13.2, WHITE, True)
        else:
            set_cell_fill(card_cell, WHITE if index % 2 else TINTS[accent])
            multiplier = index
            run = paragraph.add_run(f"{table_number} × {multiplier} = {table_number * multiplier}")
            set_run_font(run, 11.7, INK, False)


def build_document():
    add_qr(QR_PATH)
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.72)
    section.bottom_margin = Cm(0.72)
    section.left_margin = Cm(0.92)
    section.right_margin = Cm(0.92)
    section.header_distance = Cm(0.3)
    section.footer_distance = Cm(0.32)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1

    header = doc.add_table(rows=1, cols=3)
    set_table_geometry(header, [3100, 6580, 1420])
    set_row_height(header.rows[0], 2.45)
    for cell in header.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    logo_paragraph = clear_paragraph(header.cell(0, 0).paragraphs[0])
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_paragraph.add_run().add_picture(str(ROOT / "assets" / "img" / "mathsgo-logo.png"), width=Cm(3.15))

    title_paragraph = clear_paragraph(header.cell(0, 1).paragraphs[0])
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.line_spacing = 0.95
    line_1 = title_paragraph.add_run(TITLE_LINE_1 + "\n")
    set_run_font(line_1, 21.5, INK, True)
    line_2 = title_paragraph.add_run(TITLE_LINE_2)
    set_run_font(line_2, 18, INK, True)

    qr_cell = header.cell(0, 2)
    qr_paragraph = clear_paragraph(qr_cell.paragraphs[0])
    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_paragraph.add_run().add_picture(str(QR_PATH), width=Cm(1.58))
    qr_label = qr_cell.add_paragraph()
    qr_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_label.paragraph_format.space_before = Pt(0)
    qr_label.paragraph_format.space_after = Pt(0)
    set_run_font(qr_label.add_run("DÉFI TABLES"), 7.7, INK, True)

    guide = doc.add_table(rows=1, cols=1)
    set_table_geometry(guide, [11100])
    set_row_height(guide.rows[0], 1.45)
    guide_cell = guide.cell(0, 0)
    guide_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(guide_cell, top=70, start=180, bottom=65, end=180)
    set_cell_fill(guide_cell, PALE_TEAL)
    set_cell_borders(guide_cell, top="A4DDD7", start="A4DDD7", bottom="A4DDD7", end="A4DDD7", size=10)
    guide_heading = clear_paragraph(guide_cell.paragraphs[0])
    set_run_font(guide_heading.add_run(GUIDE_TITLE), 10.4, TEAL, True)
    guide_text = guide_cell.add_paragraph()
    guide_text.paragraph_format.space_before = Pt(1)
    guide_text.paragraph_format.space_after = Pt(0)
    set_run_font(guide_text.add_run(GUIDE_TEXT), 9.7, INK, False)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.25
    set_run_font(spacer.add_run(" "), 3)

    grid = doc.add_table(rows=3, cols=3)
    set_table_geometry(grid, [3700, 3700, 3700])
    accents = [TEAL, BLUE, ORANGE]
    for row in grid.rows:
        set_row_height(row, 6.74)
    for index, table_number in enumerate(TABLES):
        row, col = divmod(index, 3)
        add_card(grid.cell(row, col), table_number, accents[col])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run(FOOTER_TEXT), 7.7, MUTED, False)

    doc.core_properties.title = "Mes tables de multiplication"
    doc.core_properties.subject = "Fiche de révision des tables de multiplication"
    doc.core_properties.author = "Gwenaël Bourgault"
    doc.core_properties.keywords = "tables, multiplication, calcul mental"
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_DOCX)
